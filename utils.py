import os
import json
import pandas as pd
import random
from termcolor import colored
from tqdm import tqdm
import re
import ast
import docx
from datetime import datetime
from transformers.utils.versions import require_version
from openai import OpenAI

require_version("openai>=1.5.0", "To fix: pip install openai>=1.5.0")
random.seed(42)


def extract_fn(text):
    matches = re.findall(r"\{.*?\}", text, re.DOTALL)
    if len(matches) > 0:
        return matches[0]
    else:
        return None


def fix_guidance(text):
    # Find the value of the "guidance" field and replace the newline characters with \n
    fixed_text = re.sub(
        r'("guidance":\s*")[^"]*', lambda m: m.group(0).replace("\n", "\\n"), text
    )
    return fixed_text


def verify_fn(result, mode=None):
    temp = extract_fn(result)
    if temp:
        try:
            temp = json.loads(temp)
            return temp
        except:
            try:
                temp = ast.literal_eval(temp)
                return temp
            except:
                ## 应对guidance内有未转义\n
                if mode == "guidance":
                    try:
                        temp = extract_fn(fix_guidance(result))
                        temp = json.loads(temp)
                        return temp
                    except:
                        return "Error"
                ## 应对最后花括号前缺少双引号
                try:
                    temp = re.sub(r"(\s*})", r'" \1', temp)
                    temp = json.loads(temp)
                    return temp
                except:
                    return "Error"
    else:
        result_temp = result
        if not result_temp.strip().startswith("{"):
            result_temp = "{" + result
        if not result_temp.strip().endswith("}"):
            result_temp = result_temp + "}"
        try:
            result_temp = json.loads(result_temp)
            return result_temp
        except:
            return "Error"


def validate_stage(identified_stage):
    if "1" in identified_stage:
        return "stage1"
    elif "2" in identified_stage:
        return "stage2"
    elif "3" in identified_stage:
        return "stage3"
    elif "4" in identified_stage:
        return "stage4"
    elif "5" in identified_stage:
        return "stage5"

    stages = {
        "stage1": ["Problem Defining", "problem defining"],
        "stage2": ["Exploration", "exploration"],
        "stage3": ["Integration", "integration"],
        "stage4": ["Resolution", "resolution"],
        "stage5": ["Feedback", "feedback"],
    }
    for stage, tags in stages.items():
        for tag in tags:
            if tag in identified_stage:
                return stage
    return False


def detect_stage(identified_stage, stage_id=None):
    if "1" in identified_stage:
        return "stage1"
    elif "2" in identified_stage:
        return "stage2"
    elif "3" in identified_stage:
        return "stage3"
    elif "4" in identified_stage:
        return "stage4"
    elif "5" in identified_stage:
        return "stage5"

    stages = {
        "stage1": ["Problem Defining", "problem defining"],
        "stage2": ["Exploration", "exploration"],
        "stage3": ["Integration", "integration"],
        "stage4": ["Resolution", "resolution"],
        "stage5": ["Feedback", "feedback"],
    }
    for stage, tags in stages.items():
        for tag in tags:
            if tag in identified_stage:
                return stage

    return stage_id  # default


def validate_result(result_verified, mode=None):
    try:
        keys = list(result_verified.keys())
    except:
        return "Error"

    for key in keys:
        result_verified[key] = str(result_verified[key]).strip()

    if mode == "intervention":
        if keys != ["intervention", "intervention explanation"]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            return result_verified
    elif mode == "stage":
        if keys != ["identified stage", "stage explanation"]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            if not validate_stage(result_verified["identified stage"]):
                return "Error"
            return result_verified
    elif mode == "guidance":
        if keys != ["issue", "guidance"]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            return result_verified
    elif mode == "student":
        if keys != [
            "message",
            "self-regulation",
            "reason for self-regulation",
            "co-regulation",
            "reason for co-regulation",
        ]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            return result_verified

    return "Error"


def set_agents(topic, person_num, student_profiles):
    agents = {}
    
    for role_id, profile in student_profiles.items():
        client = OpenAI(
            base_url="https://aihubmix.com/v1",
            api_key="sk-4CFE7AOiCwVbiRXw65A86b969f204a98B4Dd5bB0D97c9b70",
        )

        agents[role_id] = {
            "client": client,
            "name": profile.split(", here is your profile")[0]
            .split("You are")[-1]
            .strip(),
            "student_prompt": profile.replace("{topic}", f'"{topic}"'),
        }
    return agents


def get_prompt():
    intervention_path = "debate_question/teacher_prompts/intervention prompt.docx"
    file = docx.Document(intervention_path)
    texts = []
    for para in file.paragraphs:
        texts.append(para.text)
    texts = [t for t in texts if t != ""]
    intervention_prompt = "\n".join(texts)

    stage_prompt_path = "debate_question/teacher_prompts/stage_prompt.docx"
    file = docx.Document(stage_prompt_path)
    texts = []
    for para in file.paragraphs:
        texts.append(para.text)
    texts = [t for t in texts if t != ""]
    stage_prompt = "\n".join(texts)

    issue_prompts_dir = "debate_question/teacher_prompts/issue_prompts"
    filelist = os.listdir(issue_prompts_dir)
    filelist = [file for file in filelist if file.endswith(".docx")]
    issue_prompts = {}
    for name in filelist:
        file_path = os.path.join(issue_prompts_dir, name)
        file = docx.Document(file_path)
        texts = []
        for para in file.paragraphs:
            texts.append(para.text)
        texts = [t for t in texts if t != ""]
        issue_prompts[name.replace(".docx", "")] = "\n".join(texts)

    return intervention_prompt, stage_prompt, issue_prompts


def set_teacher():
    client = OpenAI(
        base_url="https://aihubmix.com/v1",
        api_key="sk-4CFE7AOiCwVbiRXw65A86b969f204a98B4Dd5bB0D97c9b70",
    )

    intervention_prompt, stage_prompt, issue_prompts = get_prompt()

    teacher_agent = {
        "client": client,
        "intervention_prompt": intervention_prompt,
        "stage_prompt": stage_prompt,
        "issue_prompts": issue_prompts,
    }
    return teacher_agent


def set_teacher_vice():
    # client = OpenAI(
    #     api_key="{}".format(os.environ.get("API_KEY", "0")),
    #     # base_url="http://localhost:{}/v1".format(os.environ.get("API_PORT", 8000)),
    #     base_url="http://172.0.161.37:{}/v1".format(
    #         os.environ.get("API_PORT", 30644)
    #     ),  ##华中师服务器
    # )
    client = OpenAI(
        base_url="https://aihubmix.com/v1",
        api_key="sk-4CFE7AOiCwVbiRXw65A86b969f204a98B4Dd5bB0D97c9b70",
    )

    intervention_prompt, stage_prompt, issue_prompts = get_prompt()

    teacher_agent = {
        "client": client,
        "intervention_prompt": intervention_prompt,
        "stage_prompt": stage_prompt,
        "issue_prompts": issue_prompts,
    }
    return teacher_agent


def select_student(student_agents, recording_df, except_students=[]):
    if len(recording_df) < 1:
        specific_set = [
            id for id in range(len(student_agents)) if id not in except_students
        ]
        role_id = random.choice(specific_set)
        return student_agents[role_id], role_id
    else:
        last_two_role = list(recording_df["role_id"])[-1:]
        role_id = random.randint(0, len(student_agents) - 1)
        # if role_id == last_two_role[0] and role_id == last_two_role[1] or role_id in except_students:
        if role_id == last_two_role[0] or role_id in except_students:
            specific_set = [
                id
                for id in range(len(student_agents))
                if id != role_id and id not in except_students
            ]
            role_id = random.choice(specific_set)
            return student_agents[role_id], role_id
        else:
            return student_agents[role_id], role_id


def integrate_conversation_student(
    sub_df, random_student_agent, welcome_words, start_i, end_i, cut_word_length=None
):
    conversation_previous = []
    for i in range(start_i, end_i):
        temp_dict = {
            f"message-{i+1}": {
                "speaker": sub_df.loc[i, "name"],
                "content": sub_df.loc[i, "response"],
            }
        }
        conversation_previous.append(temp_dict)
        if (i + 1) % 10 == 0:
            conversation_previous.append(
                {
                    "system clock": f"It is round {i+1} now. If possible, please speed up the discussion."
                }
            )

    if conversation_previous == []:
        conversation_previous.append({"welcome words": welcome_words})

    conversation_previous = {
        "Conversation History": conversation_previous,
    }
    conversation_previous = json.dumps(
        conversation_previous, indent=2, ensure_ascii=False
    )

    student_prompt = random_student_agent["student_prompt"]
    history = str(conversation_previous)
    prompt = student_prompt.replace("{conversation history}", history)
    content_word_length = len(prompt)
    df_start_i = start_i

    if content_word_length > cut_word_length:
        new_start_i = start_i + 1
        prompt, df_start_i = integrate_conversation_student(
            sub_df,
            random_student_agent,
            welcome_words,
            new_start_i,
            end_i,
            cut_word_length,
        )

    return prompt, df_start_i


def integrate_conversation_teacher(
    sub_df,
    teacher_agent,
    start_i,
    end_i,
    cut_word_length=None,
    mode=None,
    intervention_result_json=None,
    stage_id=None,
    stage_result_json=None,
):
    conversation_previous = []
    conversation_now = []
    for i in range(start_i, end_i):
        if i < end_i - 1:  # 历史对话
            speaker = sub_df.loc[i, "name"]
            if speaker == "teacher":
                temp = {
                    "speaker": sub_df.loc[i, "name"],
                    "intervention": sub_df.loc[i, "intervention"],
                    "intervention reason": sub_df.loc[i, "intervention explanation"],
                    "identified stage": sub_df.loc[i, "stage"],
                    "stage reason": sub_df.loc[i, "stage explanation"],
                    "identified issue": sub_df.loc[i, "issue"],
                    "guidance": sub_df.loc[i, "response"],
                }
            else:
                temp = {
                    "speaker": sub_df.loc[i, "name"],
                    "content": sub_df.loc[i, "response"],
                }

            temp_temp = {}
            for k, v in temp.items():
                if v != "":
                    temp_temp[k] = v
            conversation_previous.append({f"message-{i + 1}": temp_temp})
        else:  # 当前对话
            if mode == "intervention":
                conversation_now.append(
                    {
                        f"message-{i + 1}": {
                            "speaker": sub_df.loc[i, "name"],
                            "content": sub_df.loc[i, "response"],
                        }
                    }
                )
            elif mode == "stage":
                conversation_now.append(
                    {
                        f"message-{i + 1}": {
                            "speaker": sub_df.loc[i, "name"],
                            "content": sub_df.loc[i, "response"],
                            "intervention": intervention_result_json["intervention"],
                            "intervention explanation": intervention_result_json[
                                "intervention explanation"
                            ],
                        }
                    }
                )
            elif mode == "guidance":
                conversation_now.append(
                    {
                        f"message-{i + 1}": {
                            "speaker": sub_df.loc[i, "name"],
                            "content": sub_df.loc[i, "response"],
                            "intervention": intervention_result_json["intervention"],
                            "intervention explanation": intervention_result_json[
                                "intervention explanation"
                            ],
                            "identified stage": stage_result_json["identified stage"],
                            "stage explanation": stage_result_json["stage explanation"],
                        }
                    }
                )

    conversation_previous = {
        "Conversation History": conversation_previous,
    }
    conversation_previous = json.dumps(
        conversation_previous, indent=2, ensure_ascii=False
    )

    conversation_now = {
        "The latest utterance": conversation_now,
    }
    conversation_now = json.dumps(conversation_now, indent=2, ensure_ascii=False)

    if mode == "intervention":
        teacher_prompt = teacher_agent["intervention_prompt"]
    elif mode == "stage":
        teacher_prompt = teacher_agent["stage_prompt"]
    elif mode == "guidance":
        teacher_prompt = teacher_agent["issue_prompts"][stage_id]

    history = f"This is the conversation history:\n{conversation_previous}\n\n"
    history += f"This is the latest utterance at this moment: {conversation_now}"
    prompt = f"System:\n{teacher_prompt}\n\nUser:\n{history}"
    content_word_length = len(prompt)

    if content_word_length > cut_word_length:
        new_start_i = start_i + 1
        if mode == "intervention":
            prompt = integrate_conversation_teacher(
                sub_df,
                teacher_agent,
                new_start_i,
                end_i,
                cut_word_length=cut_word_length,
                mode=mode,
            )
        elif mode == "stage":
            prompt = integrate_conversation_teacher(
                sub_df,
                teacher_agent,
                new_start_i,
                end_i,
                cut_word_length=cut_word_length,
                mode=mode,
                intervention_result_json=intervention_result_json,
            )
        elif mode == "guidance":
            prompt = integrate_conversation_teacher(
                sub_df,
                teacher_agent,
                new_start_i,
                end_i,
                cut_word_length=cut_word_length,
                mode=mode,
                intervention_result_json=intervention_result_json,
                stage_id=stage_id,
                stage_result_json=stage_result_json,
            )

    return prompt


def student_calling(student_name, student_agent, content, mode="student", time=0):
    messages = [{"role": "user", "content": content}]

    # stream = student_agent.chat.completions.create(messages=messages, model="test",stream=True,max_tokens=500)

    stream = student_agent.chat.completions.create(
        model="aihubmix-Llama-3-1-70B-Instruct",
        messages=messages,
        max_tokens=500,
        stream=True,
    )

    result = []
    for chunk in stream:
        if len(chunk.choices) > 0:
            result.append(chunk.choices[0].delta.content)
    result = [res for res in result if isinstance(res, str)]
    result_raw = "".join(result)

    # 获取dict内容
    result_verified = verify_fn(result_raw)

    # 严格要求输出json的keys
    result_verified = validate_result(result_verified, mode=mode)

    role_error = True
    if result_verified == "Error" or not result_verified:
        add_time = time + 1
        if add_time == 2:
            print(
                f"#####################  Recalling {mode} {student_name}... Time: {add_time} --{role_error}-- --RETURN--  ####################"
            )
            return None, None, role_error
        print(
            f"#####################  Recalling {mode} {student_name}... Time: {add_time} --{role_error}--  ####################"
        )
        content_temp = (
            content
            + "\n"
            + """Note that your last response did not complete the task. Please be sure to output it in the json format with only five keys ("message","self-regulation", "reason for self-regulation" ,"co-regulation", "reason for co-regulation"):
                    {
                    "message": "your words",
                    "self-regulation": "I1, I2,I3 ... I13",
                    "reason for self-regulation": "why these self-regulation options?",
                    "co-regulation": "G1, G2, G3 ... G13",
                    "reason for co-regulation": "why these co-regulation options?"
                    }"""
        )
        result_verified, result_raw, error_record = student_calling(
            student_name, student_agent, content_temp, mode=mode, time=add_time
        )
    else:
        ## 输出格式正确后，确保每个values都是str
        error_record = False
        keys = list(result_verified.keys())
        for key in keys:
            result_verified[key] = str(result_verified[key]).strip()

    return result_verified, result_raw, error_record


# vice_teacher = set_teacher_vice()
def teacher_calling(teacher_client, content, mode="teacher", time=0):
    messages = [{"role": "user", "content": content}]

    stream = teacher_client.chat.completions.create(
        model="aihubmix-Llama-3-1-70B-Instruct",
        messages=messages,
        max_tokens=500,
        stream=True,
    )
    # if time >= 3:
    #     print('Using vice teacher')
    #     try:
    #         stream = vice_teacher['client'].chat.completions.create(messages=messages, model="test", stream=True,
    #                                                                 max_tokens=1024)
    #     except:
    #         pass

    result = []
    for chunk in stream:
        if len(chunk.choices) > 0:
            result.append(chunk.choices[0].delta.content)
    result = [res for res in result if isinstance(res, str)]
    result_raw = "".join(result)

    # 获取dict内容
    result_verified = verify_fn(result_raw)

    # 严格要求输出json的keys
    result_verified = validate_result(result_verified, mode=mode)

    return_flag = True
    if result_verified == "Error" or not result_verified:
        add_time = time + 1
        if add_time == 5:
            return None, None, return_flag
        print(
            f"#####################  Recalling teacher {mode}... Time: {add_time}  ####################"
        )
        if mode == "intervention":
            content_temp = (
                content
                + "\n"
                + """Your answer in the last round did not complete the task. Please be sure to output it in the required format without any extra description. The output format is json with two keys (intervention and intervention explanation):
                    {
                    "intervention": "yes" or "no",
                    "intervention explanation": "why this decision?"
                    }"""
            )
            result_verified, result_raw, return_flag = teacher_calling(
                teacher_client, content_temp, mode=mode, time=add_time
            )
        elif mode == "stage":
            content_temp = (
                content
                + "\n"
                + """Your answer in the last round did not complete the task. Please be sure to output it in the required format without any extra description. The output format is json with two keys (identified stage and explanation):
                                {
                                "identified stage": "## Stage X. XXX ##",
                                "stage explanation": "why this stage?"
                                }"""
            )
            result_verified, result_raw, return_flag = teacher_calling(
                teacher_client, content_temp, mode=mode, time=add_time
            )
        elif mode == "guidance":
            content_temp = (
                content
                + "\n"
                + """Your answer in the last round did not complete the task. Please be sure to output it in the required format without any extra description. The output format is json with two keys (issue and guidance):
                                       {
                                    "issue": "** X.X XXX **",
                                    "guidance": "XXX"
                                    }"""
            )
            result_verified, result_raw, return_flag = teacher_calling(
                teacher_client, content_temp, mode=mode, time=add_time
            )
    else:
        ## 输出格式正确后，确保每个values都是str
        keys = list(result_verified.keys())
        for key in keys:
            result_verified[key] = str(result_verified[key]).strip()

        if mode == "intervention":
            if result_verified["intervention"] not in ["yes", "no"]:
                if "no" in result_verified["intervention"]:
                    result_verified["intervention"] = "no"
                else:
                    result_verified["intervention"] = "yes"

        return_flag = False

    return result_verified, result_raw, return_flag


def student_record(
    recording_df,
    room_id,
    topic,
    person_num,
    random_student_agent,
    role_id,
    student_prompt,
    student_response_json,
    student_response_raw,
):
    current_time = datetime.now()
    formatted_date = current_time.strftime("%Y/%m/%d")
    formatted_time = current_time.strftime("%H:%M:%S")

    row = len(recording_df)
    recording_df.at[row, "room_id"] = int(room_id)
    recording_df.at[row, "topic"] = topic
    recording_df.at[row, "person_num"] = int(person_num)
    recording_df.at[row, "date"] = formatted_date
    recording_df.at[row, "time"] = formatted_time
    recording_df.at[row, "role_id"] = int(role_id)
    recording_df.at[row, "name"] = random_student_agent["name"]
    recording_df.at[row, "received_information"] = student_prompt
    recording_df.at[row, "response"] = student_response_json["message"]
    recording_df.at[row, "self-regulation"] = student_response_json["self-regulation"]
    recording_df.at[row, "reason for self-regulation"] = student_response_json[
        "reason for self-regulation"
    ]
    recording_df.at[row, "co-regulation"] = student_response_json["co-regulation"]
    recording_df.at[row, "reason for co-regulation"] = student_response_json[
        "reason for co-regulation"
    ]
    recording_df.at[row, "student_response_raw"] = student_response_raw

    ending_flag = False
    if "we have finished the discussion" in student_response_json["message"].lower():
        ending_flag = True

    return ending_flag


def teacher_record(
    recording_df,
    room_id,
    topic,
    person_num,
    intervention_result_json,
    intervention_raw,
    stage_result_json=None,
    stage_raw=None,
    guidance_result_json=None,
    guidance_raw=None,
    mode=None,
):
    current_time = datetime.now()
    formatted_date = current_time.strftime("%Y/%m/%d")
    formatted_time = current_time.strftime("%H:%M:%S")

    if mode == "intervention yes":
        row = len(recording_df)
        recording_df.at[row, "room_id"] = int(room_id)
        recording_df.at[row, "topic"] = topic
        recording_df.at[row, "person_num"] = int(person_num)
        recording_df.at[row, "date"] = formatted_date
        recording_df.at[row, "time"] = formatted_time
        recording_df.at[row, "role_id"] = "T"
        recording_df.at[row, "name"] = "teacher"
        recording_df.at[row, "received_information"] = ""
        recording_df.at[row, "response"] = guidance_result_json["guidance"]
        recording_df.at[row, "intervention"] = intervention_result_json["intervention"]
        recording_df.at[row, "intervention explanation"] = intervention_result_json[
            "intervention explanation"
        ]
        recording_df.at[row, "result1_raw"] = intervention_raw
        recording_df.at[row, "stage"] = stage_result_json["identified stage"]
        recording_df.at[row, "stage explanation"] = stage_result_json[
            "stage explanation"
        ]
        recording_df.at[row, "result2_raw"] = stage_raw
        recording_df.at[row, "issue"] = guidance_result_json["issue"]
        recording_df.at[row, "guidance"] = guidance_result_json["guidance"]
        recording_df.at[row, "result3_raw"] = guidance_raw
    elif mode == "intervention no":
        row = len(recording_df) - 1
        recording_df.at[row, "intervention"] = intervention_result_json["intervention"]
        recording_df.at[row, "intervention explanation"] = intervention_result_json[
            "intervention explanation"
        ]
        recording_df.at[row, "result1_raw"] = intervention_raw


def get_student_prompt(student_name, student_profile, topic, history):
    student_prompt = f"""
    You are {student_name}, here is your profile:
    {student_profile}
    Your task:
    ##<Participating in a Group Discussion>##
    Now, you are participating in a group discussion with several participants and a teacher. The topic of the discussion is {topic}. 
    1.Your prior attitude towards this topic is attack.
    2.Your prior knowledge about this topic: [].
    3.Your status on the current topic and discussion is mainly determined by the attitude, focus, curiosity and interest in your profile. The entire discussion process should try to follow the five stages of problem definition, exploration, integration, resolution and feedback.
    4.The person with the same name as you in the conversation history is you. 
    This is the conversation history:
    {{"Conversation History":{history}}}
    ##</Participating in a Group Discussion>##
    ##<Monitoring Your Status>##
    Please monitor your self-regulation and co-regulation status at this moment. The following are the options. Any status can have only one option or multiple options.
    When I am engaged in the learning process as an individual: SELF-REGULATION
    I1: I am aware of my effort 
    I2: I am aware of my thinking 
    I3: I know my level of motivation
    I4: I question my thoughts 
    I5: I make judgments about the difficulty of a problem
    I6: I am aware of my existing knowledge
    I7: I assess my understanding 
    I8: I change my strategy when I need to
    I9: I am aware of my level of learning
    I10: I search for new strategies when needed
    I11: I apply strategies
    I12: I assess how I approach the problem
    I13: I assess my strategies 
    When I am engaged in the learning process as a member of a group: CO-REGULATION
    G1: I pay attention to the ideas of others
    G2: I listen to the comments of others
    G3: I consider the feedback of others
    G4: I reflect upon the comments of others
    G5: I observe the strategies of others
    G6: I observe how others are doing
    G7: I look for confirmation of my understanding from others
    G8: I request information from others
    G9: I respond to the contributions that others make
    G10: I challenge the strategies of others
    G11: I challenge the perspectives of others
    G12: I help the learning of others 
    G13: I monitor the learning of other
    ##</Monitoring Your Status>##
    ##<Speaking or Ending>##
    1.Please express your opinion in a sentence of no more than 100 words. 
    2.Note that your personal profile and the statuses of self-regulation and co-regulation at this moment will affect your words and speaking style, but there is no need to present these reasons in your words.
    3.Your words are usually said to the whole group. Of course, you can @ a specific participant when you want to say something to him/her, but this only happens occasionally.
    4.At the end of the discussion, you can discuss with each other that whether you can end the conversation until anyone of you says "we have finished the discussion".
    5.You will actively describe and discuss a point in depth, such as explaining its meaning and providing specific solutions.
    6.You must consider the guidance provided by your teacher and follow his or her direction to adjust and elaborate your views and words.
    7.You tend to use examples to state your views and answer the teacher's questions as much as possible. You will not always ask others what they do or think, nor will you just express superficially whether you agree with others' views.
    8.You don't have to say what you think. You just need to communicate your thoughts directly as if you were chatting normally, explain your point of view, and provide your solution.
    However, you are required to follow the following json format with only five keys ("message","self-regulation","reason for self-regulation","co-regulation","reason for co-regulation"):
    {{
    "message": "your words",
    "self-regulation": "I1, I2,I3 ... I13",
    "reason for self-regulation": "why these self-regulation options?",
    "co-regulation": "G1, G2, G3 ... G13",
    "reason for co-regulation": "why these co-regulation options?",
    }}
    ##</Speaking or Ending>##
    """
    return student_prompt
