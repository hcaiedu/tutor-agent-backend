#!/usr/bin/env python
# -*- encoding: utf-8 -*-
"""
@filename: agent_utils.py
@description:
@time: 2025/06/24
@author: Haoli WANG
@Version: 1.0
"""

import os
import pandas as pd
import random
from termcolor import colored
from tqdm import tqdm
import docx
from datetime import datetime
from transformers.utils.versions import require_version
from openai import OpenAI
from prompt_utils import extract_fn, verify_fn, fix_guidance
from flask import jsonify


def validate_stage(identified_stage):
    if "1" in identified_stage:
        return "stage1"
    elif "2" in identified_stage:
        return "stage2"
    elif "3" in identified_stage:
        return "stage3"
    elif "4" in identified_stage:
        return "stage4"
    elif "5" in identified_stage:
        return "stage5"

    stages = {
        "stage1": ["Problem Defining", "problem defining"],
        "stage2": ["Exploration", "exploration"],
        "stage3": ["Integration", "integration"],
        "stage4": ["Resolution", "resolution"],
        "stage5": ["Feedback", "feedback"],
    }
    for stage, tags in stages.items():
        for tag in tags:
            if tag in identified_stage:
                return stage
    return False


def validate_result(result_verified, mode=None):
    try:
        keys = list(result_verified.keys())
    except:
        return "Error"

    for key in keys:
        result_verified[key] = str(result_verified[key]).strip()

    if mode == "intervention":
        if keys != ["intervention", "intervention explanation"]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            return result_verified
    elif mode == "stage":
        if keys != ["identified stage", "stage explanation"]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            if not validate_stage(result_verified["identified stage"]):
                return "Error"
            return result_verified
    elif mode == "guidance":
        if keys != ["issue", "guidance"]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            return result_verified
    elif mode == "student":
        if keys != [
            "message",
            "self-regulation",
            "reason for self-regulation",
            "co-regulation",
            "reason for co-regulation",
        ]:
            return "Error"
        else:
            for key in keys:
                if result_verified[key] == "":
                    return "Error"
            return result_verified

    return "Error"


def construct_student_prompt(
    student_name,
    student_profile,
    topic,
    history,
    start_i,
    cut_word_length,
):
    """
    Description:
        use to construct student prompt
    Parameters:
        student_name(str): student name
        student_profile(json/dict): some basic profile of students
        topic(str): the dicussion topic of this room
        history(list): the conversation history of this dicussion room
        start_i(int): Determine the position from which the agent should start reviewing historical conversations
        cut_word_length(int): max length of the prompt
    Returns:
        student_prompt(str): the constructed student prompt
        dp_start_i(int): new start postion flag
    Outputs:
        None
    """
    # construct conversation history
    conversation_previous = []
    for i in range(start_i, len(history)):
        temp_dict = {
            f"message-{i+1}": {
                "speaker": history[i]["userName"],
                "content": history[i]["response"],
            }
        }
        conversation_previous.append(temp_dict)
        if (i + 1) % 10 == 0:
            conversation_previous.append(
                {
                    "system clock": f"It is round {i+1} now. If possible, please speed up the discussion."
                }
            )

    # integrate student prompt
    student_prompt = f"""
    You are {student_name}, here is your profile:
    {student_profile}
    Your task:
    ##<Participating in a Group Discussion>##
    Now, you are participating in a group discussion with several participants and a teacher. The topic of the discussion is {topic}. 
    1.Your prior attitude towards this topic is attack.
    2.Your prior knowledge about this topic: [].
    3.Your status on the current topic and discussion is mainly determined by the attitude, focus, curiosity and interest in your profile. The entire discussion process should try to follow the five stages of problem definition, exploration, integration, resolution and feedback.
    4.The person with the same name as you in the conversation history is you. 
    This is the conversation history:
    {{"Conversation History":{conversation_previous}}}
    ##</Participating in a Group Discussion>##
    ##<Monitoring Your Status>##
    Please monitor your self-regulation and co-regulation status at this moment. The following are the options. Any status can have only one option or multiple options.
    When I am engaged in the learning process as an individual: SELF-REGULATION
    I1: I am aware of my effort 
    I2: I am aware of my thinking 
    I3: I know my level of motivation
    I4: I question my thoughts 
    I5: I make judgments about the difficulty of a problem
    I6: I am aware of my existing knowledge
    I7: I assess my understanding 
    I8: I change my strategy when I need to
    I9: I am aware of my level of learning
    I10: I search for new strategies when needed
    I11: I apply strategies
    I12: I assess how I approach the problem
    I13: I assess my strategies 
    When I am engaged in the learning process as a member of a group: CO-REGULATION
    G1: I pay attention to the ideas of others
    G2: I listen to the comments of others
    G3: I consider the feedback of others
    G4: I reflect upon the comments of others
    G5: I observe the strategies of others
    G6: I observe how others are doing
    G7: I look for confirmation of my understanding from others
    G8: I request information from others
    G9: I respond to the contributions that others make
    G10: I challenge the strategies of others
    G11: I challenge the perspectives of others
    G12: I help the learning of others 
    G13: I monitor the learning of other
    ##</Monitoring Your Status>##
    ##<Speaking or Ending>##
    1.Please express your opinion in a sentence of no more than 100 words. 
    2.Note that your personal profile and the statuses of self-regulation and co-regulation at this moment will affect your words and speaking style, but there is no need to present these reasons in your words.
    3.Your words are usually said to the whole group. Of course, you can @ a specific participant when you want to say something to him/her, but this only happens occasionally.
    4.At the end of the discussion, you can discuss with each other that whether you can end the conversation until anyone of you says "we have finished the discussion".
    5.You will actively describe and discuss a point in depth, such as explaining its meaning and providing specific solutions.
    6.You must consider the guidance provided by your teacher and follow his or her direction to adjust and elaborate your views and words.
    7.You tend to use examples to state your views and answer the teacher's questions as much as possible. You will not always ask others what they do or think, nor will you just express superficially whether you agree with others' views.
    8.You don't have to say what you think. You just need to communicate your thoughts directly as if you were chatting normally, explain your point of view, and provide your solution.
    However, you are required to follow the following json format with only five keys ("message","self-regulation","reason for self-regulation","co-regulation","reason for co-regulation"):
    {{
    "message": "your words",
    "self-regulation": "I1, I2,I3 ... I13",
    "reason for self-regulation": "why these self-regulation options?",
    "co-regulation": "G1, G2, G3 ... G13",
    "reason for co-regulation": "why these co-regulation options?",
    }}
    ##</Speaking or Ending>##
    """
    content_word_length = len(student_prompt)
    df_start_i = start_i
    # check the prompt length
    if content_word_length > cut_word_length:
        new_start_i = start_i + 1
        student_prompt, df_start_i = construct_student_prompt(
            student_name,
            student_profile,
            topic,
            history,
            new_start_i,
            cut_word_length,
        )
    return student_prompt, df_start_i


def set_student_agent(client, studentInfo, roomInfo, cut_word_length):
    """
    Description:
        set student agent
    Parameters:
        client(Obj): client created by OpenAI mode
        studentInfo(dict): student info (id and name)
        roomInfo(dict): room info
        start_i(int): Determine the position from which the agent should start reviewing historical conversations
        welcome_words(str): words used at the beginning of discussion
        cut_word_length(int): max length of the prompt
        user_collection(mondodb collection): the user collection, use to get more student info
    Returns:
        student_agent(Obj): the constructed student agent
        new_start_i(int): new start postion flag
    Outputs:
        output1(type): desc
    """
    # get prompt
    student_prompt, new_start_i = construct_student_prompt(
        studentInfo["userName"],
        studentInfo["profile"],
        roomInfo["topic"],
        roomInfo["history"],
        roomInfo["start_i"],
        cut_word_length,
    )
    student_agent = {
        "client": client,
        "name": studentInfo["userName"],
        "student_prompt": student_prompt,
    }
    return student_agent, new_start_i


def student_calling(student_agent, model_name, mode="student", time=0):
    client = student_agent["client"]
    name = student_agent["name"]
    content = student_agent["student_prompt"]
    messages = [{"role": "user", "content": content}]
    stream = student_agent["client"].chat.completions.create(
        model=model_name,
        messages=messages,
        max_tokens=500,
        stream=True,
    )

    result = []
    for chunk in stream:
        if len(chunk.choices) > 0:
            result.append(chunk.choices[0].delta.content)
    result = [res for res in result if isinstance(res, str)]
    result_raw = "".join(result)

    # 获取dict内容
    result_verified = verify_fn(result_raw)

    # 严格要求输出json的keys
    result_verified = validate_result(result_verified, mode=mode)

    role_error = True
    if result_verified == "Error" or not result_verified:
        add_time = time + 1
        if add_time == 2:
            print(
                f"#####################  Recalling {mode} {name}... Time: {add_time} --{role_error}-- --RETURN--  ####################"
            )
            return None, None, role_error
        print(
            f"#####################  Recalling {mode} {name}... Time: {add_time} --{role_error}--  ####################"
        )
        content_temp = (
            content
            + "\n"
            + """Note that your last response did not complete the task. Please be sure to output it in the json format with only five keys ("message","self-regulation", "reason for self-regulation" ,"co-regulation", "reason for co-regulation"):
                    {
                    "message": "your words",
                    "self-regulation": "I1, I2,I3 ... I13",
                    "reason for self-regulation": "why these self-regulation options?",
                    "co-regulation": "G1, G2, G3 ... G13",
                    "reason for co-regulation": "why these co-regulation options?"
                    }"""
        )
        result_verified, result_raw, error_record = student_calling(
            student_agent["name"], student_agent, content_temp, mode=mode, time=add_time
        )
    else:
        ## 输出格式正确后，确保每个values都是str
        error_record = False
        keys = list(result_verified.keys())
        for key in keys:
            result_verified[key] = str(result_verified[key]).strip()

    return result_verified, result_raw, error_record


def get_student_agent_response(
    client,studentInfo, roomInfo, model_name, cut_word_length=18000
):
    """
    Description:

    Parameters:
      param1(type): desc
    Returns:
      return1(type): desc
    Outputs:
      output1(type): desc
    """
    # set student agent
    student_agent, start_i = set_student_agent(client, studentInfo, roomInfo, cut_word_length)
    try:
        # call student agent api
        student_response_json, student_response_raw, role_error = student_calling(
            student_agent,
            model_name,
            mode="student",
            time=0,
        )
        # handle role_error
        if role_error:
            return "role_error", str(studentInfo["userId"])
        else:
            return "success", student_response_json
    except Exception as e:
        return "error", "system error"
      
def set_teacher_prompt(
    student_name,
    student_response,
    room_info,
    teacher_agent,
    cut_word_length=None,
    mode=None,
    intervention_result_json=None,
    stage_id=None,
    stage_result_json=None,
):
    conversation_previous = []
    conversation_now = []
    history = room_info["history"]
    start_i = room_info["start_i"]
    # handle prvevious conversation
    for i in range(
        start_i, len(history)
    ):  # this history is from last query of database, so it doesn't have the text that teacher need to determine whether intervent or not
        speaker = history[i]["userName"]
        if speaker == "teacher":
            temp = {
                "speaker": speaker,
                "intervention": history[i]["intervention"],
                "intervention reason": history[i]["intervention explanation"],
                "identified stage": history[i]["stage"],
                "stage reason": history[i]["stage explanation"],
                "identified issue": history[i]["issue"],
                "guidance": history[i]["response"],
            }
        else:
            temp = {
                "speaker": history[i]["userName"],
                "content": history[i]["response"],
            }
        conversation_previous.append({f"message-{i + 1}": temp})
    # handle current text (student response)
    if mode == "intervention":
        conversation_now.append(
            {
                f"message-{ len(history) + 1}": {
                    "speaker": student_name,
                    "content": student_response,
                }
            }
        )
    elif mode == "stage":
        conversation_now.append(
            {
                f"message-{ len(history) + 1}": {
                    "speaker": student_name,
                    "content": student_response,
                    "intervention": intervention_result_json["intervention"],
                    "intervention explanation": intervention_result_json[
                        "intervention explanation"
                    ],
                }
            }
        )
    elif mode == "guidance":
        conversation_now.append(
            {
                f"message-{ len(history) + 1}": {
                    "speaker": student_name,
                    "content": student_response,
                    "intervention": intervention_result_json["intervention"],
                    "intervention explanation": intervention_result_json[
                        "intervention explanation"
                    ],
                    "identified stage": stage_result_json["identified stage"],
                    "stage explanation": stage_result_json["stage explanation"],
                }
            }
        )
    if mode == "intervention":
        teacher_prompt = teacher_agent["intervention_prompt"]
    elif mode == "stage":
        teacher_prompt = teacher_agent["stage_prompt"]
    elif mode == "guidance":
        teacher_prompt = teacher_agent["issue_prompts"][stage_id]

    prompt = f"""
        System:
        {teacher_prompt}
        User:
        This is the conversation history:
        {{"Conversation History":{conversation_previous}}}
        This is the latest utterance at this moment: 
        {{"The latest utterance": {conversation_now}}}
    """

    content_word_length = len(prompt)

    if content_word_length > cut_word_length:
        new_start_i = start_i + 1
        if mode == "intervention":
            prompt = set_teacher_prompt(
                student_name,
                student_response,
                room_info,
                teacher_agent,
                cut_word_length=cut_word_length,
                mode=mode,
            )
        elif mode == "stage":
            prompt = set_teacher_prompt(
                student_name,
                student_response,
                room_info,
                teacher_agent,
                cut_word_length=cut_word_length,
                mode=mode,
                intervention_result_json=intervention_result_json,
            )
        elif mode == "guidance":
            prompt = set_teacher_prompt(
                student_name,
                student_response,
                room_info,
                teacher_agent,
                cut_word_length=cut_word_length,
                mode=mode,
                intervention_result_json=intervention_result_json,
                stage_id=stage_id,
                stage_result_json=stage_result_json,
            )

    return prompt
  
def teacher_calling(teacher_client, content, model_name, mode="teacher", time=0):
    messages = [{"role": "user", "content": content}]

    stream = teacher_client.chat.completions.create(
        model=model_name,
        messages=messages,
        max_tokens=500,
        stream=True,
    )

    result = []
    for chunk in stream:
        if len(chunk.choices) > 0:
            result.append(chunk.choices[0].delta.content)
    result = [res for res in result if isinstance(res, str)]
    result_raw = "".join(result)

    # 获取dict内容
    result_verified = verify_fn(result_raw)

    # 严格要求输出json的keys
    result_verified = validate_result(result_verified, mode=mode)

    return_flag = True
    if result_verified == "Error" or not result_verified:
        add_time = time + 1
        if add_time == 5:
            return None, None, return_flag
        print(
            f"#####################  Recalling teacher {mode}... Time: {add_time}  ####################"
        )
        if mode == "intervention":
            content_temp = (
                content
                + "\n"
                + """Your answer in the last round did not complete the task. Please be sure to output it in the required format without any extra description. The output format is json with two keys (intervention and intervention explanation):
                    {
                    "intervention": "yes" or "no",
                    "intervention explanation": "why this decision?"
                    }"""
            )
            result_verified, result_raw, return_flag = teacher_calling(
                teacher_client, content_temp, model_name, mode=mode, time=add_time
            )
        elif mode == "stage":
            content_temp = (
                content
                + "\n"
                + """Your answer in the last round did not complete the task. Please be sure to output it in the required format without any extra description. The output format is json with two keys (identified stage and explanation):
                                {
                                "identified stage": "## Stage X. XXX ##",
                                "stage explanation": "why this stage?"
                                }"""
            )
            result_verified, result_raw, return_flag = teacher_calling(
                teacher_client, content_temp, model_name, mode=mode, time=add_time
            )
        elif mode == "guidance":
            content_temp = (
                content
                + "\n"
                + """Your answer in the last round did not complete the task. Please be sure to output it in the required format without any extra description. The output format is json with two keys (issue and guidance):
                                       {
                                    "issue": "** X.X XXX **",
                                    "guidance": "XXX"
                                    }"""
            )
            result_verified, result_raw, return_flag = teacher_calling(
                teacher_client, content_temp, model_name, mode=mode, time=add_time
            )
    else:
        ## 输出格式正确后，确保每个values都是str
        keys = list(result_verified.keys())
        for key in keys:
            result_verified[key] = str(result_verified[key]).strip()

        if mode == "intervention":
            if result_verified["intervention"] not in ["yes", "no"]:
                if "no" in result_verified["intervention"]:
                    result_verified["intervention"] = "no"
                else:
                    result_verified["intervention"] = "yes"

        return_flag = False

    return result_verified, result_raw, return_flag
  
def get_prompt():
    """
    Description:
        Function used to get the teacher prompts
    Parameters:
        None
    Returns:
        intervention_prompt(str): consturcted intervention prompt
        stage_prompt(str): consturcted stage prompt
        issue_prompts(dict): consturcted issue prompts
    Outputs:
        None
    """
    intervention_path = "prompts/teacher_prompts/intervention prompt.docx"
    file = docx.Document(intervention_path)
    texts = []
    for para in file.paragraphs:
        texts.append(para.text)
    texts = [t for t in texts if t != ""]
    intervention_prompt = "\n".join(texts)
    stage_prompt_path = "prompts/teacher_prompts/stage_prompt.docx"
    file = docx.Document(stage_prompt_path)
    texts = []
    for para in file.paragraphs:
        texts.append(para.text)
    texts = [t for t in texts if t != ""]
    stage_prompt = "\n".join(texts)
    issue_prompts_dir = "prompts/teacher_prompts/issue_prompts"
    filelist = os.listdir(issue_prompts_dir)
    filelist = [file for file in filelist if file.endswith(".docx")]
    issue_prompts = {}
    for name in filelist:
        file_path = os.path.join(issue_prompts_dir, name)
        file = docx.Document(file_path)
        texts = []
        for para in file.paragraphs:
            texts.append(para.text)
        texts = [t for t in texts if t != ""]
        issue_prompts[name.replace(".docx", "")] = "\n".join(texts)
    return intervention_prompt, stage_prompt, issue_prompts


def set_teacher(client):
    """
    Description:
        set teacher agent
    Parameters:
        client(Obj): client created by OpenAI mode
    Returns:
        teacher_agent(Obj): and object which contains one openai mode client and three teacher prompts
    Outputs:
        None
    """
    intervention_prompt, stage_prompt, issue_prompts = get_prompt()
    teacher_agent = {
        "client": client,
        "intervention_prompt": intervention_prompt,
        "stage_prompt": stage_prompt,
        "issue_prompts": issue_prompts,
    }
    return teacher_agent
  
def detect_stage(identified_stage, stage_id=None):
    if "1" in identified_stage:
        return "stage1"
    elif "2" in identified_stage:
        return "stage2"
    elif "3" in identified_stage:
        return "stage3"
    elif "4" in identified_stage:
        return "stage4"
    elif "5" in identified_stage:
        return "stage5"

    stages = {
        "stage1": ["Problem Defining", "problem defining"],
        "stage2": ["Exploration", "exploration"],
        "stage3": ["Integration", "integration"],
        "stage4": ["Resolution", "resolution"],
        "stage5": ["Feedback", "feedback"],
    }
    for stage, tags in stages.items():
        for tag in tags:
            if tag in identified_stage:
                return stage

    return stage_id
  
def get_teacher_response(
    client, student_name, student_response_json, room_info, model_name, cut_word_length=18000
):
    """
    Description:

    Parameters:
      param1(type): desc
    Returns:
      return1(type): desc
    Outputs:
      output1(type): desc
    """
    # set teacher agent
    teacher_agent = set_teacher(client)
    # teacher intervention
    intervention_prompt = set_teacher_prompt(
        student_name,
        student_response_json,
        room_info,
        teacher_agent,
        cut_word_length=18000,
        mode="intervention",
    )
    intervention_result_json, intervention_raw, return_flag = teacher_calling(
        teacher_agent["client"], intervention_prompt, model_name, mode="intervention", time=0
    )
    intervention_flag = intervention_result_json["intervention"]
    # if intervention is needed
    if intervention_flag == "yes":
        stage_prompt = set_teacher_prompt(
            student_name,
            student_response_json,
            room_info,
            teacher_agent,
            cut_word_length=18000,
            mode="stage",
            intervention_result_json=intervention_result_json,
        )
        stage_result_json, stage_raw, return_flag = teacher_calling(
            teacher_agent["client"], stage_prompt, model_name, mode="stage", time=0
        )
        stage_id = detect_stage(stage_result_json["identified stage"], stage_id=stage_id)
        guidance_prompt = set_teacher_prompt(
            student_name,
            student_response_json,
            room_info,
            teacher_agent,
            cut_word_length=18000,
            mode="guidance",
            intervention_result_json=intervention_result_json,
            stage_id=stage_id,
            stage_result_json=stage_result_json,
        )
        guidance_result_json, guidance_raw, return_flag = teacher_calling(
            teacher_agent["client"], guidance_prompt, model_name, mode="guidance", time=0
        )
        # add teacher response to history
        current_time = datetime.now()  # get date and time
        formatted_date = current_time.strftime("%Y/%m/%d")
        formatted_time = current_time.strftime("%H:%M:%S")
        history_item = {
            "date": formatted_date,
            "time": formatted_time,
            "userId": "T",
            "userName": "teacher",
            "userAvatar": "../assets/Agent.PNG",
            "received_information": "",
            "response": guidance_result_json["guidance"],
            "intervention": intervention_result_json["intervention"],
            "intervention explanation": intervention_result_json[
                    "intervention explanation"
                ],
            "result1_raw": intervention_raw,
            "stage": stage_result_json["identified stage"],
            "stage explanation": stage_result_json[
                    "stage explanation"
                ],
            "result2_raw": stage_raw,
            "issue": guidance_result_json["issue"],
            "guidance": guidance_result_json["guidance"],
            "result3_raw": guidance_raw
        }
        return "intervention_yes", history_item
        # append_result = roomcollection.update_one(
        #     {"_id": roomInfo["_id"]},
        #     {"$push": {"history": history_item}}
        # )
        # turn_counting += 1
    else:
        # no intervention is needed, add intervention result to student conversation history
        message = student_response_json
        message["intervention"] = intervention_result_json["intervention"]
        message["result1_raw"] = intervention_raw
        message["intervention explanation"] = intervention_result_json["intervention explanation"]
        return "intervention_no", message

